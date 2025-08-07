import pandas as pd
import os
import shutil
import mysql.connector
import re
import glob
from conn import config
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from docx2pdf import convert

# === ARQUIVOS E PASTAS ===
pasta_raiz = "Certidoes_Emitidas"
template_filename = 'Certidao_sem_dados.docx'
assinatura_imagem = 'Campos_Assinatura.jpg'
os.makedirs(pasta_raiz, exist_ok=True)

# === FORMATAR CPF ===
def formatar_cpf(cpf: str) -> str:
    cpf = ''.join(filter(str.isdigit, str(cpf)))
    return f"{cpf[:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}" if len(cpf) == 11 else cpf

# === FORMATAÇÃO TABELA ===
def adicionar_borda_em_tabela(table):
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for borda in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{borda}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def reduzir_margens_tabela(table, margem_em_twips=50):
    tbl = table._element
    tblPr = tbl.tblPr
    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ('top', 'start', 'bottom', 'end'):
        cell_margin = OxmlElement(f'w:{side}')
        cell_margin.set(qn('w:w'), str(margem_em_twips))
        cell_margin.set(qn('w:type'), 'dxa')
        tblCellMar.append(cell_margin)
    tblPr.append(tblCellMar)

# === CONSULTAS AO BANCO ===
def buscar_atividades(id_aluno, id_curso):
    try:
        conn = mysql.connector.connect(**config)
        cursor = conn.cursor()
        query = """
            SELECT gi.itemname AS nome_atividade, g.finalgrade AS nota
            FROM mdl_grade_items gi
            JOIN mdl_grade_grades g ON g.itemid = gi.id
            WHERE gi.courseid = %s AND g.userid = %s AND gi.itemtype <> 'course' AND g.finalgrade IS NOT NULL
            ORDER BY g.timemodified;
        """
        cursor.execute(query, (id_curso, id_aluno))
        return cursor.fetchall()
    except:
        return []
    finally:
        if conn.is_connected():
            cursor.close()
            conn.close()

def buscar_nota_final(id_aluno, id_curso):
    try:
        conn = mysql.connector.connect(**config)
        cursor = conn.cursor()
        query = """
            SELECT g.finalgrade
            FROM mdl_grade_items gi
            JOIN mdl_grade_grades g ON g.itemid = gi.id
            WHERE gi.courseid = %s AND g.userid = %s AND gi.itemtype = 'course';
        """
        cursor.execute(query, (id_curso, id_aluno))
        resultado = cursor.fetchone()
        return resultado[0] if resultado else None
    except:
        return None
    finally:
        if conn.is_connected():
            cursor.close()
            conn.close()

# === SUBSTITUIR MARCADORES COM NEGRITO ===
def substituir_marcadores_formatados(paragraph, substitutions):
    full_text = ''.join(run.text for run in paragraph.runs)
    if not any(marcador in full_text for marcador in substitutions):
        return
    for i in range(len(paragraph.runs)):
        paragraph.runs[i].text = ''
    partes = re.split(r'({{.*?}})', full_text)
    for parte in partes:
        if parte in substitutions:
            run = paragraph.add_run(substitutions[parte])
            run.bold = True
        else:
            paragraph.add_run(parte)

# === BUSCAR E PROCESSAR PLANILHAS ===
planilhas = glob.glob(os.path.join(pasta_raiz, "*.xlsx"))

if not planilhas:
    print("⚠️ Nenhuma planilha .xlsx encontrada em Certidoes_Emitidas.")
else:
    for planilha in planilhas:
        try:
            print(f"\n📥 Processando: {planilha}")
            df = pd.read_excel(planilha)

            if df.empty:
                print(f"⚠️ Planilha vazia: {planilha}")
                continue

            curso_nome = df.iloc[0]["curso"]
            curso_id = df.iloc[0]["id_curso"]
            pasta_saida = os.path.join(pasta_raiz, f"{curso_nome} [{curso_id}]")
            os.makedirs(pasta_saida, exist_ok=True)

            for _, row in df.iterrows():
                try:
                    nome = row['nome_completo']
                    cpf = formatar_cpf(row['username'])
                    id_aluno = row['id_aluno']
                    id_curso = row['id_curso']

                    print(f"📄 Gerando certidão para: {nome}...")

                    doc = Document(template_filename)

                    for p in doc.paragraphs:
                        substituir_marcadores_formatados(p, {
                            "{{nome}}": nome,
                            "{{cpf}}": cpf,
                            "{{curso}}": curso_nome
                        })

                    doc.add_page_break()

                    titulo = doc.add_paragraph()
                    run = titulo.add_run("Histórico de Notas do Curso")
                    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.font.size = Pt(14)
                    run.bold = True

                    table = doc.add_table(rows=1, cols=2)
                    adicionar_borda_em_tabela(table)
                    reduzir_margens_tabela(table)
                    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.columns[0].width = Inches(4.85)
                    table.columns[1].width = Inches(0.15)

                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    run0 = hdr_cells[0].paragraphs[0].add_run("Atividade")
                    run0.bold = True
                    run0.font.size = Pt(8)

                    run1 = hdr_cells[1].paragraphs[0].add_run("Nota")
                    run1.bold = True
                    run1.font.size = Pt(8)

                    atividades = buscar_atividades(id_aluno, id_curso)
                    for atividade, nota in atividades:
                        row_cells = table.add_row().cells
                        row_cells[0].text = atividade
                        row_cells[1].text = f"{nota:.2f}"
                        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(7)
                                paragraph.paragraph_format.space_after = Pt(0)

                    nota_final = buscar_nota_final(id_aluno, id_curso)
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    run = p.add_run(f"Nota Final do Curso: {nota_final:.2f}" if nota_final is not None else "")
                    run.font.size = Pt(9)
                    run.bold = True

                    if os.path.exists(assinatura_imagem):
                        par = doc.add_paragraph()
                        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        par.paragraph_format.space_before = Pt(0)
                        par.paragraph_format.space_after = Pt(0)
                        run = par.add_run()
                        run.add_picture(assinatura_imagem, width=Inches(5.2))

                    docx_nome = f"Certidao - {nome}.docx"
                    docx_path = os.path.join(pasta_saida, docx_nome)
                    doc.save(docx_path)
                    convert(docx_path)
                    os.remove(docx_path)

                    print(f"✅ PDF gerado com sucesso para: {nome}")

                except Exception as erro:
                    print(f"❌ Erro ao gerar para {row.get('nome_completo', '---')}: {erro}")

            # Move a planilha para a pasta do curso
            shutil.move(planilha, os.path.join(pasta_saida, os.path.basename(planilha)))
            print(f"📁 Planilha movida para: {pasta_saida}")

        except Exception as e:
            print(f"❌ Erro ao processar '{planilha}': {e}")
